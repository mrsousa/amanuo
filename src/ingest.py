"""Gera contexto (.md) e dados fixos (contracts.json) a partir de DOCUMENTOS.

Fonte mais rica que os emails: proposta comercial, atas, contrato. Você
organiza os arquivos por contrato e o Claude destila **cada documento
individualmente** em duas saídas: campos estruturados + contexto em
Markdown. Processar documento por documento (em vez de juntar tudo antes de
enviar) evita que um arquivo novo fique de fora por causa do teto de
caracteres — o corte, se acontecer, afeta só aquele documento.

A pasta raiz onde ficam esses documentos é configurável (chave "fontes_dir"
em config.json, padrão "fontes") para não colidir com a pasta docs/ do
próprio repositório. As pastas podem ter o nome completo do projeto — o ID
do contrato é extraído automaticamente do início do nome:

    fontes/
      DF21-100 - Vale - Projeto Conceitual e Detalhado/
        proposta.pdf
        ata_2024-03.docx
      DF22-C038 - Vale - Projeto Detalhado Maravilhas III/
        contrato.pdf

Padrão do código (as duas letras identificam a unidade de negócio, ex. DF =
DF+ Engenharia; "C" no padrão novo indica contrato):
    antigo: AA00-000   (ex.: DF21-100)
    novo:   AA00-C000  (ex.: DF22-C038)

O código extraído (DF21-100, DF22-C038) é o ID usado em contracts.json e em
contexts/<id>/ — precisa bater com o ID gerado por 'run.py init' a partir
das pastas do Outlook.

O contexto de cada documento fica em contexts/<id>/<arquivo>.md, e
contexts/<id>.md é a junção de todos eles (regenerada a cada rodada, sem
custo de API — é só concatenação de arquivo).

Uso:
    python ingest.py                        # todos os contratos, todos os documentos
    python ingest.py DF22-C038               # só esse contrato, todos os documentos
    python ingest.py DF22-C038 proposta.pdf  # só esse documento desse contrato
"""

import json
import os
import sys
import glob

import classifier
from graph_client import load_config

MAX_CHARS = 20000  # teto de texto por DOCUMENTO enviado ao modelo

INGEST_ROLE = (
    "Você lê um documento de contrato por vez (proposta, ata, contrato) e "
    "produz: (1) os dados fixos estruturados e (2) um contexto em Markdown "
    "que servirá para uma IA decidir, no futuro, se novos emails pertencem "
    "a este contrato. Extraia apenas o que estiver neste documento; não "
    "invente e não repita o que provavelmente já veio de outro documento."
)


# ---------- extração de texto ----------
def extract_pdf(path: str) -> str:
    import pdfplumber
    partes = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            partes.append(page.extract_text() or "")
    return "\n".join(partes)


def extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    linhas = [p.text for p in doc.paragraphs]
    # também pega texto de tabelas (comum em propostas)
    for tbl in doc.tables:
        for row in tbl.rows:
            linhas.append(" | ".join(c.text for c in row.cells))
    return "\n".join(linhas)


def extract_text(path: str) -> str | None:
    """Extrai o texto de um PDF/DOCX. None se a extensão não for suportada."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".docx", ".doc"):
        return extract_docx(path)
    return None


# ---------- destilação via Claude ----------
def distill(config, cid: str, nome_arquivo: str, texto: str) -> tuple[dict, str, dict]:
    """Retorna (campos_fixos, contexto_md, usage) para UM documento."""
    user = (
        f"Documento \"{nome_arquivo}\" do contrato \"{cid}\":\n\n{texto[:MAX_CHARS]}\n\n"
        "Produza SOMENTE um JSON com duas chaves:\n"
        '{"fields": {"cliente": "", "numero_contrato": "", "objeto": "", '
        '"valor": "", "data_inicio": "", "prazo": "", "responsavel": ""}, '
        '"context_md": "<contexto em Markdown com o que ESTE documento traz: '
        'partes, pessoas-chave, escopo, marcos, pendências e vocabulário '
        'típico; ~150 palavras. Omita categorias sem informação aqui>"}\n'
        "Deixe em branco os campos que não encontrar. Não invente dados."
    )
    msg = classifier._client(config).messages.create(
        model=config.get("model", "claude-sonnet-5"),
        max_tokens=1500,
        system=[{"type": "text", "text": INGEST_ROLE}],
        messages=[
            {"role": "user", "content": user},
            {"role": "assistant", "content": "{"},  # prefill: força JSON
        ],
    )
    usage = classifier._usage_dict(msg.usage)
    try:
        data = json.loads("{" + msg.content[0].text)
    except json.JSONDecodeError:
        return {}, "", usage
    fields = {k: v for k, v in (data.get("fields") or {}).items() if v}
    return fields, data.get("context_md", ""), usage


# ---------- orquestração ----------
def _recombinar_contexto(cid: str, rec: dict) -> None:
    """Junta contexts/<id>/<arquivo>.md em contexts/<id>.md (sem custo de
    API — só concatenação). Roda depois de processar cada documento, então
    o arquivo combinado nunca fica dessincronizado dos documentos individuais."""
    pasta_doc = os.path.join(classifier.CONTEXTS_DIR, cid)
    if not os.path.isdir(pasta_doc):
        return
    titulo = rec.get("folder_name") or rec.get("cliente") or cid
    blocos = []
    for nome in sorted(os.listdir(pasta_doc)):
        conteudo = open(os.path.join(pasta_doc, nome), encoding="utf-8").read().strip()
        blocos.append(f"### {nome[:-3]}\n\n{conteudo}")  # remove sufixo .md do nome exibido
    with open(os.path.join(classifier.CONTEXTS_DIR, f"{cid}.md"), "w", encoding="utf-8") as f:
        f.write(f"# {titulo} ({cid})\n\n" + "\n\n".join(blocos) + "\n")


def process_document(config, cid: str, path: str, contracts: dict) -> dict | None:
    """Processa UM documento: destila, atualiza contracts.json e grava
    contexts/<id>/<arquivo>.md (recombinando o contexts/<id>.md consolidado)."""
    nome = os.path.basename(path)
    try:
        texto = extract_text(path)
    except Exception as e:  # não deixa um arquivo ruim derrubar a rodada
        print(f"    ! falha lendo {nome}: {e}")
        return None
    if texto is None:
        return None  # extensão não suportada, silenciosamente ignorada
    if not texto.strip():
        print(f"  {cid}/{nome}: sem texto extraível, pulando")
        return None

    print(f"  {cid}/{nome}: destilando {len(texto)} chars...")
    fields, context_md, usage = distill(config, cid, nome, texto)

    # --- dados fixos → contracts.json (preserva metadados de pasta do init) ---
    rec = contracts.get(cid, {})
    preservar = {k: rec[k] for k in ("folder_id", "folder_name", "folder_path", "ativo") if k in rec}
    rec.update(fields)
    rec.update(preservar)
    contracts[cid] = rec

    # --- contexto deste documento → contexts/<id>/<arquivo>.md ---
    if context_md:
        pasta_doc = os.path.join(classifier.CONTEXTS_DIR, cid)
        os.makedirs(pasta_doc, exist_ok=True)
        with open(os.path.join(pasta_doc, f"{nome}.md"), "w", encoding="utf-8") as f:
            f.write(context_md.strip() + "\n")
        _recombinar_contexto(cid, rec)
    return usage


def main():
    config = load_config()
    contracts = classifier.load_contracts()
    fontes_dir = config.get("fontes_dir", "fontes")

    if not os.path.isdir(fontes_dir):
        sys.exit(f"Pasta '{fontes_dir}/' não existe. Crie {fontes_dir}/<contrato>/ com os documentos.")
    todas_pastas = [d for d in os.listdir(fontes_dir)
                    if os.path.isdir(os.path.join(fontes_dir, d))]

    alvo_cid = sys.argv[1] if len(sys.argv) > 1 else None
    alvo_arquivo = sys.argv[2] if len(sys.argv) > 2 else None

    if alvo_cid:
        pastas = [d for d in todas_pastas if classifier.extract_cid(d) == alvo_cid]
        if not pastas:
            sys.exit(f"Nenhuma pasta em {fontes_dir}/ corresponde ao contrato '{alvo_cid}'.")
    else:
        pastas = todas_pastas
    if not pastas:
        sys.exit(f"Nenhum contrato encontrado em {fontes_dir}/.")

    total = {"input": 0, "output": 0, "cache_read": 0, "cache_write": 0}
    for folder_name in sorted(pastas):
        cid = classifier.extract_cid(folder_name)
        folder = os.path.join(fontes_dir, folder_name)
        arquivos = sorted(glob.glob(os.path.join(folder, "*")))

        if alvo_arquivo:
            arquivos = [p for p in arquivos if os.path.basename(p) == alvo_arquivo]
            if not arquivos:
                sys.exit(f"Arquivo '{alvo_arquivo}' não encontrado em {folder}.")

        for path in arquivos:
            u = process_document(config, cid, path, contracts)
            if u:
                for k in total:
                    total[k] += u[k]

    classifier.save_contracts(contracts)
    print(f"\ncontracts.json e {classifier.CONTEXTS_DIR}/ atualizados.")
    print(f"[tokens: entrada {total['input']}, saída {total['output']}]")


if __name__ == "__main__":
    main()
